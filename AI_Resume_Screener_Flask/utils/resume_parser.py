"""
utils/resume_parser.py
Extracts structured information from resume PDF and DOCX files.
Uses pdfplumber for PDF and python-docx for DOCX.
"""

import re
import io
import pdfplumber
from docx import Document


# ── Text Extraction ─────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file (provided as bytes)."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        text = f"[PDF parsing error: {e}]"
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file (provided as bytes)."""
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as e:
        text = f"[DOCX parsing error: {e}]"
    return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Auto-detect file type and extract text."""
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif fname.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    return ""


# ── Field Extraction ─────────────────────────────────────────────────────────

def extract_name(text: str) -> str:
    """Heuristic: the name is usually in the first few lines."""
    lines = [l.strip() for l in text.strip().splitlines() if l.strip()]
    for line in lines[:6]:
        # A name usually has 2-4 words, no special characters
        if 2 <= len(line.split()) <= 4 and not any(c in line for c in ["@", ".", "+", "/", ":"]):
            return line
    return lines[0] if lines else "Unknown"


def extract_email(text: str) -> str | None:
    """Extract the first email address found in the text."""
    match = re.search(r"[\w.+\-]+@[\w\-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else None


def extract_phone(text: str) -> str | None:
    """Extract a phone number using common formats."""
    match = re.search(r"(\+?\d[\d\s\-().]{7,14}\d)", text)
    return match.group(0).strip() if match else None


def extract_linkedin(text: str) -> str | None:
    """Extract LinkedIn profile URL."""
    match = re.search(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    return "https://" + match.group(0) if match else None


def extract_github(text: str) -> str | None:
    """Extract GitHub profile URL."""
    match = re.search(r"github\.com/[\w\-]+", text, re.IGNORECASE)
    return "https://" + match.group(0) if match else None


def extract_portfolio(text: str) -> str | None:
    """Extract portfolio URL (common patterns)."""
    match = re.search(r"https?://(?!linkedin|github)[\w\-\.]+\.[a-z]{2,}(?:/[\w\-\./?=&]*)?", text, re.IGNORECASE)
    return match.group(0) if match else None


def extract_years_experience(text: str) -> float:
    """
    Estimate years of experience.
    First tries explicit mentions like '3 years experience',
    then counts employment date ranges.
    """
    # Try explicit mention
    patterns = [
        r"(\d+)\+?\s+years?\s+(?:of\s+)?(?:total\s+)?experience",
        r"experience\s+of\s+(\d+)\+?\s+years?",
        r"(\d+)\+?\s+years?\s+(?:in|working)",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return float(m.group(1))

    # Count date ranges like "2021 – 2024" or "2022 – Present"
    ranges = re.findall(
        r"(20\d{2})\s*[-–]\s*(20\d{2}|present|current|now)",
        text, re.IGNORECASE
    )
    total_years = 0.0
    for start, end in ranges:
        s = int(start)
        e = 2025 if end.lower() in ("present", "current", "now") else int(end)
        total_years += max(0, e - s)
    return min(total_years, 25.0)


def extract_education(text: str) -> list:
    """Extract education degrees mentioned in the text."""
    degree_keywords = {
        "PhD": ["ph.d", "phd", "doctorate"],
        "M.Tech": ["m.tech", "mtech", "m.e.", "m.e "],
        "M.Sc": ["m.sc", "msc", "master of science"],
        "MBA": ["mba", "master of business"],
        "B.Tech": ["b.tech", "btech", "b.e.", "b.e "],
        "B.Sc": ["b.sc", "bsc", "bachelor of science"],
        "BCA": ["bca"],
        "MCA": ["mca"],
        "Diploma": ["diploma"],
    }
    text_lower = text.lower()
    found = []
    for degree, keywords in degree_keywords.items():
        if any(kw in text_lower for kw in keywords):
            found.append(degree)
    return found


def extract_certifications(text: str) -> list:
    """Extract certification keywords from the resume text."""
    cert_keywords = [
        "aws certified", "google cloud certified", "azure certified",
        "tensorflow developer", "deep learning specialization",
        "pmp", "scrum master", "cissp", "ceh",
        "coursera", "udemy", "edx", "nptel", "hackerrank",
    ]
    text_lower = text.lower()
    return [c.title() for c in cert_keywords if c in text_lower]


# ── Main Parser ──────────────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Main resume parsing function.
    Returns a dictionary of extracted fields.
    """
    raw_text = extract_text(file_bytes, filename)

    return {
        "raw_text": raw_text,
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "linkedin": extract_linkedin(raw_text),
        "github": extract_github(raw_text),
        "portfolio": extract_portfolio(raw_text),
        "years_experience": extract_years_experience(raw_text),
        "education": extract_education(raw_text),
        "certifications": extract_certifications(raw_text),
    }
